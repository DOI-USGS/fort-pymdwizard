#!/usr/bin/env python
# -*- coding: utf8 -*-
"""
The MetadataWizard (pymdwizard) software was developed by the U.S. Geological
Survey Fort Collins Science Center.

License:            Creative Commons Attribution 4.0 International (CC BY 4.0)
                    https://creativecommons.org/licenses/by/4.0/

PURPOSE
------------------------------------------------------------------------------
Module contains functionality for creating a metadata review document.


NOTES
------------------------------------------------------------------------------
None
"""

# Standard python libraries.
import os
import time
import getpass

# Non-standard python libraries.
try:
    from docx import Document
    from docx.shared import Inches
    from docx.shared import Pt
    from docx.shared import RGBColor
    from docx.enum.style import WD_STYLE_TYPE
except ImportError as err:
    raise ImportError(err, __file__)

# Custom import/libraries.
try:
    from pymdwizard.core import (utils, fgdc_utils)
except ImportError as err:
    raise ImportError(err, __file__)

# Global variable that loads the local resource 'bdp_lookup' into a JSON
# object.
FGDC_LOOKUP = fgdc_utils.get_fgdc_lookup()


def _get_longname(tag):
    """
    Description:
        Retrieve the long name associated with a given tag from the FGDC
        lookup dictionary. If the tag is not found, return the tag itself.

    Args:
        tag (str): The tag for which the long name is to be fetched.

    Returns:
        str: The corresponding long name if found, otherwise the original tag.
    """

    try:
        # Try to get the long name from the FGDC lookup.
        long_name = FGDC_LOOKUP[tag]["long_name"]
    except KeyError:
        # If the tag is not found, use the tag as the long name.
        long_name = tag

    return long_name


def _load_styles(doc):
    """
    Description:
        Load and define custom styles for a Word document.

    Args:
        doc (Document): The document object to which styles will be added.

    Returns:
        None
    """

    # Set the default 'Normal' style.
    heading1 = doc.styles["Normal"]
    font = heading1.font
    font.name = "Times New Roman"
    font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    font.size = Pt(24)
    font.bold = True

    # Define custom styles for FGDC title and headings.
    styles = doc.styles
    new_heading_style = styles.add_style("fgdc title", WD_STYLE_TYPE.PARAGRAPH)
    new_heading_style.base_style = styles["Heading 1"]
    font = new_heading_style.font
    font.name = "Times New Roman"
    font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    font.size = Pt(24)
    font.bold = True

    # FGDC heading 2 style.
    new_heading_style = \
        styles.add_style("fgdc heading 2", WD_STYLE_TYPE.PARAGRAPH)
    new_heading_style.base_style = styles["Heading 2"]
    font = new_heading_style.font
    font.name = "Times New Roman"
    font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    font.size = Pt(18)
    font.bold = True

    # FGDC heading 3 style.
    new_heading_style = \
        styles.add_style("fgdc heading 3", WD_STYLE_TYPE.PARAGRAPH)
    new_heading_style.base_style = styles["Heading 3"]
    font = new_heading_style.font
    font.name = "Times New Roman"
    font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    font.size = Pt(13.5)
    font.bold = True

    # Review content heading style.
    new_heading_style = styles.add_style(
        "review content heading", WD_STYLE_TYPE.PARAGRAPH
    )
    new_heading_style.base_style = styles["Heading 3"]
    font = new_heading_style.font
    font.name = "Times New Roman"
    font.color.rgb = RGBColor(0x36, 0x5F, 0x91)
    font.size = Pt(13.5)
    font.bold = True

    # FGDC bar style.
    new_heading_style = \
        styles.add_style("fgdc bar", WD_STYLE_TYPE.PARAGRAPH)
    new_heading_style.base_style = styles["Normal"]
    font = new_heading_style.font
    font.name = "Times New Roman"
    font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    font.size = Pt(12)
    font.bold = True

    # FGDC tag style.
    new_heading_style = styles.add_style("fgdc tag", WD_STYLE_TYPE.PARAGRAPH)
    new_heading_style.base_style = styles["Normal"]
    font = new_heading_style.font
    font.name = "Times New Roman"
    font.color.rgb = RGBColor(0x48, 0x8A, 0xC7)
    font.size = Pt(12)
    font.italic = True

    # FGDC tag content paragraph style.
    new_heading_style = \
        styles.add_style("fgdc tag content p", WD_STYLE_TYPE.PARAGRAPH)
    new_heading_style.base_style = styles["Normal"]
    font = new_heading_style.font
    font.name = "Times New Roman"
    font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    font.size = Pt(12)
    font.italic = False
    font.bold = False

    # FGDC tag content character style.
    new_heading_style = \
        styles.add_style("fgdc tag content c", WD_STYLE_TYPE.CHARACTER)
    new_heading_style.base_style = styles["Normal"]
    font = new_heading_style.font
    font.name = "Times New Roman"
    font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    font.size = Pt(12)
    font.italic = True
    font.bold = True

    # FGDC tag content problem character style.
    new_heading_style = \
        styles.add_style("fgdc tag content problem c", WD_STYLE_TYPE.CHARACTER)
    new_heading_style.base_style = styles["Normal"]
    font = new_heading_style.font
    font.name = "Times New Roman"
    font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    font.size = Pt(12)
    font.italic = True
    font.bold = False

    # FGDC bold style.
    new_heading_style = styles.add_style("fgdc bold", WD_STYLE_TYPE.PARAGRAPH)
    new_heading_style.base_style = styles["Normal"]
    font = new_heading_style.font
    font.name = "Times New Roman"
    font.color.rgb = RGBColor(0x15, 0x15, 0x15)
    font.size = Pt(12)
    font.italic = False
    font.bold = True

    # FGDC link style.
    new_heading_style = styles.add_style("fgdc link", WD_STYLE_TYPE.PARAGRAPH)
    new_heading_style.base_style = styles["List Bullet"]
    font = new_heading_style.font
    font.name = "Times New Roman"
    font.color.rgb = RGBColor(0x00, 0x00, 0xFF)
    font.size = Pt(12)
    font.underline = True
    font.bold = False


def _add_tag(
    doc,
    tag,
    content="",
    indent=0,
    tag_style="fgdc tag",
    content_style="fgdc tag content",
):
    """
    Description:
        Adds a tag to the document with optional content.

    Args:
        doc (Document): The document object to which the tag and content will
            be added.
        tag (str): The tag text to add.
        content (str, optional): The content to add next to the tag; defaults
            to an empty string.
        indent (int, optional): The indent level for the content; defaults to 0.
        tag_style (str, optional): The style to apply to the tag; defaults to
            "fgdc tag".
        content_style (str, optional): The style to apply to the content;
            defaults to "fgdc tag content".

    Returns:
        cit (Paragraph): The created paragraph object for the tag.
    """

    # Add the tag as a paragraph with the specified style.
    cit = doc.add_paragraph(tag + ": ", style=tag_style)
    cit.paragraph_format.space_after = Inches(0.005)

    if content:
        if len(content) > 70:
            content = doc.add_paragraph(content, style=content_style + " p")
            content.paragraph_format.left_indent = Inches(indent)
            content.paragraph_format.line_spacing = 1
        else:
            content = cit.add_run(content)
            content.style = content_style + " c"

    return cit


def _add_child_content(doc, node, indent=0.25):
    """
    Description:
        Adds child content to a document based on the provided node,
        including handling for onlink nodes.

    Args:
        doc (Document): The document object to which content will be added.
        node (Element): The XML node whose content and children will be
            processed.
        indent (float, optional): The indent level for child content; defaults
            to 0.25 inches.

    Returns:
        None
    """

    # Add the main line for the current node with its tag and text content.
    line = _add_tag(doc, _get_longname(node.tag), node.text, indent)

    # Check if the node is an 'onlink' type and verify the URL.
    if node.tag == "onlink":
        if not utils.url_is_alive(node.text):
            line2 = line.add_run(" (check url, possible problem)",
                                 style='fgdc tag content problem c')

    # Set the left indent for the current line.
    line.paragraph_format.left_indent = Inches(indent)

    # Recursively add child content for each child node.
    for child in node.children:
        _add_child_content(doc, child, indent + 0.25)


def format_errors(xml_document, which="bdp"):
    """
    Description:
        Formats errors in the given XML document based on the specified schema.

    Args:
        xml_document (Document): The XML document that contains the errors.

        which (str, optional): pecifies which schema to use for validation.
            Defaults to "bdp".

    Returns:
        validation_result (bool): The result of the validation process.
    """

    # Determine the appropriate XSL file based on the specified schema.
    if which == "bdp":
        xsl_fname = \
            utils.get_resource_path("FGDC/BDPfgdc-std-001-1998-annotated.xsd")
    else:
        xsl_fname = \
            utils.get_resource_path("FGDC/fgdc-std-001-1998-annotated.xsd")

    # Validate the XML document against the selected XSL schema.
    return fgdc_utils.validate_xml(xml_document.record, xsl_fname)


def generate_review_report(xml_document, docx_fname, which="bdp"):
    """
    Description:
        Generates a review report based on the provided XML document
        and saves it as a DOCX file.

    Args:
        xml_document (Document): The XML document containing metadata to review.
        docx_fname (str): The filename for the output DOCX report.
        which (str, optional): The type of schema to validate against; defaults
            to "bdp".

    Returns:
        None
    """

    # Create a new Word Document.
    document = Document()

    # Set document proofing properties.
    DOCX = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    element = document.settings.element.find(DOCX + "proofState")
    element.attrib[DOCX + "grammar"] = "dirty"
    element.attrib[DOCX + "spelling"] = "dirty"

    # Load the predefined styles for the document.
    _load_styles(document)

    # Prepare the title of the report.
    title_str = "Metadata Review for:\n\t{}"
    try:
        title_str = xml_document.metadata.idinfo.citation.citeinfo.title.text
    except:
        title_str = (
            "No FGDC Title found in record at "
            "metadta/idinfo/citation/citeinfo/title"
        )

    # Add title to the document.
    title = document.add_heading("Metadata Review for:", level=2)
    title.style = document.styles["fgdc heading 2"]
    title.paragraph_format.line_spacing = 1
    title2 = document.add_paragraph("\t" + title_str + "\n",
                                    style="fgdc tag content p")

    # Add a visual separator to the title.
    bar = title2.add_run("_" * 72)
    document.add_paragraph("")

    # Add review information section.
    reviewinfoline = document.add_heading("Review Info:", level=3)
    reviewinfoline.style = document.styles["review content heading"]
    fname = os.path.split(xml_document.fname)[-1]

    # Add the filename being reviewed.
    _add_tag(document, "Metadata file being reviewed", fname,
             tag_style="fgdc bold")

    # Retrieve reviewer information.
    try:
        username = getpass.getuser()
        contact = utils.get_usgs_contact_info(username, True)
        reviewer_str = "{} ({})".format(
            contact["fgdc_cntperp"]["fgdc_cntper"],
            contact["fgdc_cntemail"]
        )
    except:
        # something when wrong getting the contact info
        # (no internet, non USGS, etc.), just insert placeholcers
        reviewer_str = "<<insert reviewer name>> (<<insert reviewer email>>)"

    # Add reviewer information.
    _add_tag(document, "Reviewer", reviewer_str, tag_style="fgdc bold")
    _add_tag(document, "Review Date", time.strftime("%m/%d/%Y"),
             tag_style="fgdc bold")
    document.add_paragraph("")

    # Add summary section.
    summaryline = document.add_heading("Summary:", level=3)
    summaryline.style = document.styles["review content heading"]
    document.add_paragraph(
        "<<insert general review comments here>>",
        style="fgdc tag content p"
    )

    # Add space for comments.
    document.add_paragraph("", style="fgdc tag content p")
    document.add_paragraph("", style="fgdc tag content p")

    # Add errors section based on the schema validation.
    errorline = document.add_heading("FGDC Schema Errors:", level=3)
    errorline.style = document.styles["review content heading"]

    # Get errors from validation.
    errors = format_errors(xml_document, which)
    if errors:
        document.add_paragraph("Error (XML Path to error)",
                               style="fgdc tag")
        for error in errors:
            e = document.add_paragraph(
                "\t{}\n\t({})".format(error[1], error[0]),
                style="fgdc tag content p"
            )
    else:
        document.add_paragraph("\tNo Schema Errors found",
                               style="fgdc tag")

    # Add space after errors section.
    document.add_paragraph("")

    # Add metadata content section.
    mdcontentline = document.add_heading("Metadata Content:", level=3)
    mdcontentline.style = document.styles["review content heading"]
    bar = document.add_paragraph("_" * 72, style="fgdc bar")

    title2 = document.add_heading("Metadata:", level=3)
    title2.style = document.styles["fgdc heading 3"]

    # Loop through and add each child metadata element.
    for child in xml_document.metadata.children:
        long_name = _get_longname(child.tag)
        link = document.add_paragraph(text=long_name, style="fgdc link")
        link.paragraph_format.left_indent = Inches(0.5)
        link.paragraph_format.line_spacing = 1

    for child in xml_document.metadata.children:
        long_name = _get_longname(child.tag)
        bar = document.add_paragraph("_" * 72, style="fgdc bar")

        section_title = document.add_heading(long_name + ":", level=3)
        section_title.style = document.styles["fgdc heading 3"]
        section_title.paragraph_format.space_after = Inches(0.15)

        # Add child content recursively
        _add_child_content(document, child)

    # Save the generated document to the provided filename.
    document.save(docx_fname)
